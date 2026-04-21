from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

C_H1    = RGBColor(0x26, 0x07, 0x00)
C_H2    = RGBColor(0xFF, 0x63, 0x2B)
C_H3    = RGBColor(0x26, 0x07, 0x00)
C_BODY  = RGBColor(0x1A, 0x1A, 0x1A)
C_NOTE  = RGBColor(0x55, 0x55, 0x55)
FONT      = 'Calibri'
FONT_CODE = 'Courier New'

CODE_STARTERS = re.compile(
    r'^(\s{4,}|#|def |for |while |if |else:|elif |import |from |return |with |class |'
    r'print\(|!pip|!apt|[a-z_]+\s*=\s|[a-z_]+\.[a-z_]+\(|'
    r'pd\.|np\.|df\.|sns\.|plt\.|drive\.|resultado|cuadrados|valores |x,\s*y|'
    r'strategy=|\.fit\(|\.transform\(|\.mount\()'
)
CODE_LABELS = re.compile(r'^(Python|R)\s*$', re.IGNORECASE)


def is_code(text):
    return bool(CODE_STARTERS.match(text)) or bool(CODE_LABELS.match(text.strip()))


def shade_paragraph(para, hex_color='F0F0F0'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_left_border(para, color='FF632B', sz='18'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bd = OxmlElement('w:left')
    bd.set(qn('w:val'), 'single')
    bd.set(qn('w:sz'), sz)
    bd.set(qn('w:space'), '100')
    bd.set(qn('w:color'), color)
    pBdr.append(bd)
    pPr.append(pBdr)


def apply_run(para, text, size, bold, color, font=FONT, italic=False):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def format_para(para, ptype, text):
    para.clear()
    pf = para.paragraph_format
    pf.left_indent = Inches(0)

    if ptype == 'h1':
        apply_run(para, text, 20, True, C_H1)
        pf.space_before = Pt(22)
        pf.space_after  = Pt(8)

    elif ptype == 'h2':
        apply_run(para, text, 15, True, C_H2)
        pf.space_before = Pt(14)
        pf.space_after  = Pt(5)

    elif ptype == 'h3':
        apply_run(para, text, 12, True, C_H3)
        pf.space_before = Pt(10)
        pf.space_after  = Pt(4)

    elif ptype == 'code':
        if not text:
            return
        apply_run(para, text, 9.5, False, RGBColor(0x1E, 0x1E, 0x1E), FONT_CODE)
        pf.left_indent   = Inches(0.3)
        pf.space_before  = Pt(2)
        pf.space_after   = Pt(2)
        shade_paragraph(para)

    elif ptype == 'note':
        apply_run(para, text, 10.5, False, C_NOTE, italic=True)
        pf.left_indent  = Inches(0.3)
        pf.space_before = Pt(4)
        pf.space_after  = Pt(6)
        add_left_border(para)

    elif ptype == 'list':
        apply_run(para, text, 11, False, C_BODY)
        pf.left_indent  = Inches(0.3)
        pf.space_before = Pt(1)
        pf.space_after  = Pt(3)

    else:  # body
        apply_run(para, text, 11, False, C_BODY)
        pf.space_before = Pt(2)
        pf.space_after  = Pt(6)


# ===================== CLASE 02 =====================
def format_clase02():
    path = 'c:/Users/Turi/Desktop/Data Science I/Data Science 1/Clase02/Clase 02.docx'
    doc = Document(path)
    in_code = False

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            in_code = False
            continue

        style = para.style.name

        if style == 'Heading 2':
            format_para(para, 'h1', txt)
            in_code = False

        elif style == 'Heading 3':
            if CODE_LABELS.match(txt):
                format_para(para, 'code', '')
                in_code = True
            elif in_code or is_code(txt):
                format_para(para, 'code', txt)
                in_code = True
            elif txt.startswith('Nota:') or txt.startswith('Importante:'):
                format_para(para, 'note', txt)
                in_code = False
            elif re.match(r'^[A-ZÁÉÍÓÚÑ¿].{0,80}$', txt) and len(txt) < 90 and not txt[0].islower():
                words = txt.split()
                # Short rules like "Indentación: usar 4 espacios..." → list item
                if re.match(r'^[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+:', txt) and len(words) > 3:
                    format_para(para, 'list', txt)
                elif len(words) <= 8:
                    format_para(para, 'h2', txt)
                else:
                    format_para(para, 'body', txt)
                in_code = False
            else:
                format_para(para, 'body', txt)
                in_code = False

    doc.save(path)
    print('Clase 02 done.')


# ===================== CLASE 04 =====================
H2_TITLES = {
    'Estructura recomendada para notebooks',
    'Gestion de dependencias en Colab',
    'Gestión de dependencias en Colab',
    'Integracion con la nube (Google Drive)',
    'Integración con la nube (Google Drive)',
    'Tipos de valores ausentes',
    'Tecnicas de deteccion con Pandas',
    'Técnicas de detección con Pandas',
    'Estrategias Iniciales con Pandas',
    'Estrategias Intermedias',
    'Estrategias Avanzadas',
    'SimpleImputer (Univariante)',
    'IterativeImputer (Multivariante)',
    'Flujo de trabajo (ColumnTransformer)',
    'Conversion a datetime',
    'Conversión a datetime',
    'Creacion de indices temporales y Slicing',
    'Creación de índices temporales y Slicing',
    '1. Introduccion',
    '1. Introducción',
}


def format_clase04():
    path = 'c:/Users/Turi/Desktop/Data Science I/Data Science 1/Clase04/Clase 04.docx'
    doc = Document(path)

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue

        # Code block: starts with "Python " or "R "
        if re.match(r'^(Python|R)\s+\S', txt):
            code_txt = re.sub(r'^(Python|R)\s+', '', txt)
            format_para(para, 'code', code_txt)

        elif re.match(r'^\d+[\.\s\-]', txt) and len(txt) < 100:
            # Only treat as H1 if it's a real section (has meaningful description after number)
            after_num = re.sub(r'^\d+[\.\s\-]+', '', txt)
            if len(after_num.split()) >= 4:
                format_para(para, 'h1', txt)
            else:
                format_para(para, 'h3', txt)

        elif txt in H2_TITLES or any(txt.startswith(t) for t in H2_TITLES):
            format_para(para, 'h2', txt)

        elif txt.startswith('Nota:') or txt.startswith('Importante:'):
            format_para(para, 'note', txt)

        elif re.match(r'^(MCAR|MAR|MNAR|Pros:|Contras:|strategy=)', txt):
            format_para(para, 'list', txt)

        elif len(txt.split()) <= 6 and re.match(r'^[A-ZÁÉÍÓÚÑ]', txt):
            format_para(para, 'h2', txt)

        else:
            format_para(para, 'body', txt)

    doc.save(path)
    print('Clase 04 done.')


format_clase02()
format_clase04()
print('All done.')
